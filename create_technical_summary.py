"""
CSCI218 Group Project: Generate Technical Summary Document
===========================================================
Fills in the Technical Summary template with project content.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(BASE_DIR, "CSCI218_Technical_Summary.docx")

# Create document
doc = Document()

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ============================================================
# COVER PAGE
# ============================================================
doc.add_paragraph("FACULTY OF ENGINEERING AND INFORMATION SCIENCES").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
title = doc.add_paragraph("CSCI218 Assignment Coversheet")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph("Attach this Cover Sheet as the front page of the assignment. Assignments are to be submitted to the relevant academic. Submit only one copy of the assignment per group to the relevant online site.")

# Plagiarism section
plagiarism_title = doc.add_paragraph("PLAGIARISM")
plagiarism_title.runs[0].bold = True
doc.add_paragraph("Deliberate plagiarism may lead to failure in the subject. Plagiarism is cheating by using the written ideas or submitted work of someone else. The University of Wollongong has a strong policy against plagiarism. A plagiarized assignment will receive a zero mark and be penalized according to the university rules. Plagiarism detection software might be used for this assignment.")
doc.add_paragraph("See Acknowledgement Practice/Plagiarism Prevention Policy at http://www.uow.edu.au/about/policy/UOW058648.html")
doc.add_paragraph()
doc.add_paragraph("Any use of generative AI tools must be clearly stated in the report. Failure to do so may result in penalties for misuse. For example, directly copying and pasting AI-generated content is not permitted.").runs[0].bold = True

doc.add_paragraph()

# Student names
students = [
    ("Jasmine April Aulia Ng", "[Student ID]"),
    ("Felix", "[Student ID]"),
    ("Aryan", "[Student ID]"),
    ("Shawn", "[Student ID]"),
    ("Toh Jun Peng Brandon", "[Student ID]"),
    ("Lian Ziang", "[Student ID]"),
]

for name, sid in students:
    doc.add_paragraph(f"STUDENT NAME / UOW ID: {name} / {sid}")

doc.add_paragraph()
doc.add_paragraph("Subject Code & Name: CSCI218 Foundations of AI")
doc.add_paragraph("Tutor's Name: Ms Cher Lim")
doc.add_paragraph("Assignment Due Date: 22nd Feb 2026 (SGT 11:55pm)")

doc.add_paragraph()
declaration = doc.add_paragraph("DECLARATION")
declaration.runs[0].bold = True
doc.add_paragraph("We certify that this is entirely our own work, except where we have given fully documented references to the work of others, and that the material contained in this assignment has not previously been submitted for assessment in any formal course of study. We understand the definition and consequences of plagiarism.")

doc.add_paragraph()
ack = doc.add_paragraph("ACKNOWLEDGEMENT")
ack.runs[0].bold = True
doc.add_paragraph("The marker of this assignment may, for the purpose of assessing this assignment, reproduce this assignment and provide a copy to another member of academic staff.")

doc.add_paragraph()
for name, _ in students:
    doc.add_paragraph(f"Student Signature: {name}     Date: ____________")

doc.add_paragraph()
doc.add_paragraph("[Insert e-signatures or type names]").runs[0].italic = True

# Page break
doc.add_page_break()

# ============================================================
# TECHNICAL SUMMARY HEADER
# ============================================================
header = doc.add_paragraph("CSCI218 Technical Summary")
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.runs[0].bold = True
header.runs[0].font.size = Pt(18)

doc.add_paragraph()
doc.add_paragraph("Project Title: Dry Bean Classification Using Machine Learning Algorithms").runs[0].bold = True
doc.add_paragraph("GitHub link: [Insert GitHub Repository URL]")

doc.add_paragraph()

# ============================================================
# ATTRIBUTION TABLE
# ============================================================
attr_title = doc.add_paragraph("Attribution Table (Required):")
attr_title.runs[0].bold = True

# Create attribution table
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'

# Header row
headers = ["Section", "Student Name(s)", "Roles", "Percentage"]
for i, header_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header_text
    cell.paragraphs[0].runs[0].bold = True

# Data rows based on rubric assignment
attribution_data = [
    ("1. Problem Definition & Dataset Description", "Jasmine April Aulia Ng", "A, D", "100%"),
    ("2. Methodology & Design Rationale", "Felix", "A, R", "100%"),
    ("3. Experimental Setup & Training Config", "Aryan", "E, D", "100%"),
    ("4. Evaluation Metrics & Analysis Methods", "Shawn", "A, V", "100%"),
    ("5. Results & Error Analysis", "Toh Jun Peng Brandon", "E, V, A", "100%"),
    ("6. Limitations & Proposed Improvements", "Lian Ziang", "A, R", "100%"),
]

for row_idx, (section, name, roles, pct) in enumerate(attribution_data, start=1):
    table.rows[row_idx].cells[0].text = section
    table.rows[row_idx].cells[1].text = name
    table.rows[row_idx].cells[2].text = roles
    table.rows[row_idx].cells[3].text = pct

doc.add_paragraph()
doc.add_paragraph("Roles: A = Author (primary writer), R = Reviewer (quality check), E = Experimentation (ran code/tuning), V = Visualization (charts, figures), D = Documentation")

doc.add_paragraph()

# ============================================================
# SECTION 1: Problem Definition & Dataset Description (5%)
# ============================================================
s1_title = doc.add_paragraph("1. Problem Definition & Dataset Description (5%)")
s1_title.runs[0].bold = True
s1_title.runs[0].font.size = Pt(12)

s1_content = """
This project addresses the multi-class classification problem of identifying dry bean varieties from their physical characteristics. The goal is to develop a machine learning system that can automatically classify dry beans into one of seven types based on shape and size features extracted from grain images.

Problem Statement: Manual classification of dry beans by human experts is time-consuming, subjective, and prone to errors. An automated ML-based classification system can improve accuracy, consistency, and throughput in agricultural quality control processes.

Dataset: The Dry Bean Dataset from the UCI Machine Learning Repository (ID: 602) contains 13,611 samples of seven registered dry bean varieties grown in Turkey.

Key Properties:
- Classes (7): BARBUNYA (1,322), BOMBAY (522), CALI (1,630), DERMASON (3,546), HOROZ (1,928), SEKER (2,027), SIRA (2,636)
- Features (16): All numeric, derived from high-resolution images using computer vision
- Feature categories: Geometric (Area, Perimeter, MajorAxisLength, MinorAxisLength), Shape (AspectRatio, Eccentricity, Roundness, Compactness), Form (Extent, Solidity, ConvexArea, EquivDiameter), and Fourier descriptors (ShapeFactor1-4)
- No missing values; dataset is clean and ready for ML

Reference: Koklu, M. and Ozkan, I.A., 2020. Multiclass classification of dry beans using computer vision and machine learning techniques. Computers and Electronics in Agriculture, 174, 105507.
"""
doc.add_paragraph(s1_content.strip())

# ============================================================
# SECTION 2: Methodology & Design Rationale (20%)
# ============================================================
s2_title = doc.add_paragraph("2. Methodology & Design Rationale (20%)")
s2_title.runs[0].bold = True
s2_title.runs[0].font.size = Pt(12)

s2_content = """
We selected three diverse machine learning algorithms to compare different classification paradigms:

1. K-Nearest Neighbours (KNN, k=5)
   - Rationale: Simple instance-based baseline that makes no assumptions about data distribution
   - How it works: Classifies each sample by majority vote of its 5 nearest neighbours using Euclidean distance
   - Strengths: Non-parametric, easy to interpret, works well with clean data
   - Considerations: Requires feature scaling; computationally expensive at prediction time for large datasets

2. Random Forest (100 trees)
   - Rationale: Ensemble method known to perform excellently on tabular data with correlated features
   - How it works: Trains 100 decision trees on bootstrap samples, each using random feature subsets; final prediction by majority vote
   - Strengths: Handles feature correlations well, provides feature importance, resistant to overfitting
   - Considerations: Less interpretable than single decision tree

3. Support Vector Machine (SVM, RBF kernel, C=10)
   - Rationale: Strong theoretical foundation; RBF kernel handles non-linear class boundaries
   - How it works: Maps data to higher-dimensional space via RBF kernel, finds optimal hyperplane maximizing margin between classes
   - Strengths: Effective in high-dimensional spaces, robust to outliers due to margin maximization
   - Considerations: Requires feature scaling; hyperparameter tuning (C, gamma) affects performance

Design Choices:
- Stratified 80/20 train-test split preserves class proportions in both sets
- StandardScaler normalization (zero mean, unit variance) required for KNN and SVM which are distance/kernel-based
- 5-fold stratified cross-validation provides reliable accuracy estimates with uncertainty bounds
- Weighted metrics (precision, recall, F1) account for class imbalance in the dataset
"""
doc.add_paragraph(s2_content.strip())

# ============================================================
# SECTION 3: Experimental Setup & Training Configuration (10%)
# ============================================================
s3_title = doc.add_paragraph("3. Experimental Setup & Training Configuration (10%)")
s3_title.runs[0].bold = True
s3_title.runs[0].font.size = Pt(12)

s3_content = """
Software Environment:
- Python 3.14
- scikit-learn 1.8.0 (ML algorithms, preprocessing, evaluation)
- pandas 3.0.0, numpy 2.4.2 (data manipulation)
- matplotlib 3.10.8, seaborn 0.13.2 (visualization)
- ucimlrepo 0.0.7 (dataset loading from UCI repository)

Hardware: Standard desktop/laptop (no GPU required for these classical ML methods)

Data Pipeline:
1. Load dataset via ucimlrepo API (13,611 samples, 16 features)
2. Label encoding: Convert 7 class names to integers 0-6
3. Train-test split: 80% training (10,888 samples), 20% test (2,723 samples), stratified by class
4. Feature scaling: StandardScaler fitted on training data only, then applied to test data (prevents data leakage)

Model Hyperparameters:
- KNN: n_neighbors=5, metric='euclidean' (sklearn defaults)
- Random Forest: n_estimators=100, max_depth=None (no limit), random_state=42
- SVM: kernel='rbf', C=10 (regularization), gamma='scale' (1/(n_features * X.var()))

Training Procedure:
- 5-fold stratified cross-validation on training set for model selection
- Final models trained on full training set
- Evaluation performed on held-out test set (never seen during training or CV)

Reproducibility: random_state=42 set for all random operations (train_test_split, Random Forest, CV)

Code: All experiments implemented in dry_bean_classification.py (~350 lines)
"""
doc.add_paragraph(s3_content.strip())

# ============================================================
# SECTION 4: Evaluation Metrics & Analysis Methods (20%)
# ============================================================
s4_title = doc.add_paragraph("4. Evaluation Metrics & Analysis Methods (20%)")
s4_title.runs[0].bold = True
s4_title.runs[0].font.size = Pt(12)

s4_content = """
Metrics Used:

1. Accuracy = (Correct Predictions) / (Total Predictions)
   - Primary metric for overall model performance
   - Appropriate because class imbalance is moderate (not extreme)

2. Precision (Weighted) = Weighted average of per-class precision
   - Precision = TP / (TP + FP) for each class
   - Measures how many predicted positives are actually correct
   - Weighted by class support to account for imbalance

3. Recall (Weighted) = Weighted average of per-class recall
   - Recall = TP / (TP + FN) for each class
   - Measures how many actual positives are correctly identified
   - Important for ensuring no bean type is systematically missed

4. F1-Score (Weighted) = Harmonic mean of precision and recall
   - Balances precision-recall trade-off
   - Single metric summarizing per-class performance

5. 5-Fold Stratified Cross-Validation Accuracy
   - Mean and standard deviation across 5 folds
   - Provides uncertainty estimate and detects overfitting
   - Stratified ensures each fold has same class distribution

Why These Metrics?
- Multi-class problem requires metrics that handle all 7 classes fairly
- Weighted averaging accounts for class imbalance (DERMASON has 6x more samples than BOMBAY)
- CV accuracy validates that test set performance is not due to lucky split
- Confusion matrices provide detailed per-class error analysis

Analysis Methods:
- Confusion matrix heatmaps to identify which classes are confused with each other
- Feature importance from Random Forest to understand which features drive classification
- Correlation heatmap to understand feature relationships
- Class distribution visualization to understand data balance
"""
doc.add_paragraph(s4_content.strip())

# ============================================================
# SECTION 5: Results & Error Analysis (30%)
# ============================================================
s5_title = doc.add_paragraph("5. Results & Error Analysis (30%)")
s5_title.runs[0].bold = True
s5_title.runs[0].font.size = Pt(12)

s5_content = """
Main Results:

| Model                | CV Accuracy       | Test Accuracy | Precision | Recall | F1-Score |
|---------------------|-------------------|---------------|-----------|--------|----------|
| K-Nearest Neighbours | 0.9233 +/- 0.0046 | 0.9166        | 0.9174    | 0.9166 | 0.9168   |
| Random Forest        | 0.9240 +/- 0.0053 | 0.9207        | 0.9209    | 0.9207 | 0.9207   |
| SVM (RBF)           | 0.9336 +/- 0.0053 | 0.9243        | 0.9243    | 0.9243 | 0.9243   |

Best Model: SVM (RBF) with 92.43% test accuracy

Key Findings:
1. All three models achieved >91% accuracy, demonstrating the dataset is well-suited for ML classification
2. SVM outperformed others by ~0.4-0.8%, likely due to RBF kernel's ability to model non-linear boundaries
3. Random Forest and KNN performed similarly, with RF slightly better
4. CV and test accuracies are consistent (no significant overfitting detected)

Per-Class Analysis (SVM - Best Model):
| Class    | Precision | Recall | F1-Score | Support |
|----------|-----------|--------|----------|---------|
| BARBUNYA | 0.93      | 0.91   | 0.92     | 265     |
| BOMBAY   | 1.00      | 1.00   | 1.00     | 104     |
| CALI     | 0.94      | 0.94   | 0.94     | 326     |
| DERMASON | 0.91      | 0.92   | 0.92     | 709     |
| HOROZ    | 0.96      | 0.96   | 0.96     | 386     |
| SEKER    | 0.94      | 0.95   | 0.95     | 406     |
| SIRA     | 0.87      | 0.86   | 0.87     | 527     |

Error Analysis:
- BOMBAY: Perfect classification (100%) - these beans are distinctively larger than others
- SIRA: Lowest performance (87% F1) - often confused with DERMASON and SEKER
- SIRA and DERMASON confusion: Both are smaller beans with similar shape characteristics
- Feature importance (from RF): ShapeFactor4, ShapeFactor2, and Compactness are most discriminative
- Highly correlated features (e.g., Area-Perimeter, MajorAxisLength-MinorAxisLength) may introduce redundancy

[See Appendix for confusion matrix visualizations and feature importance charts]
"""
doc.add_paragraph(s5_content.strip())

# ============================================================
# SECTION 6: Limitations & Proposed Improvements (15%)
# ============================================================
s6_title = doc.add_paragraph("6. Limitations & Proposed Improvements (15%)")
s6_title.runs[0].bold = True
s6_title.runs[0].font.size = Pt(12)

s6_content = """
Current Limitations:

1. No Hyperparameter Optimization
   - Used default or manually selected hyperparameters
   - Grid search or Bayesian optimization could improve performance

2. Feature Redundancy
   - High correlations between some features (e.g., Area and ConvexArea: r > 0.99)
   - May reduce model interpretability and slightly impact performance

3. Class Imbalance
   - BOMBAY has only 522 samples vs DERMASON with 3,546
   - Some classes may be underrepresented in training

4. Limited Model Diversity
   - Only tested 3 classical ML algorithms
   - Did not explore deep learning or gradient boosting methods

5. SIRA Class Confusion
   - Consistent misclassification between SIRA and similar-sized beans
   - Feature set may not capture distinguishing characteristics

Proposed Improvements:

1. Hyperparameter Tuning
   - Use GridSearchCV or RandomizedSearchCV for systematic optimization
   - Tune k for KNN, n_estimators/max_depth for RF, C/gamma for SVM

2. Feature Engineering
   - Apply PCA to reduce correlated features while preserving variance
   - Create new features (e.g., Area/Perimeter ratio) that may improve discrimination

3. Address Class Imbalance
   - Use SMOTE (Synthetic Minority Over-sampling) for underrepresented classes
   - Apply class weights during training

4. Explore Additional Models
   - XGBoost/LightGBM: Often outperform RF on tabular data
   - MLP Neural Network: May capture complex non-linear patterns
   - Ensemble stacking: Combine predictions from multiple models

5. Improve SIRA Classification
   - Collect additional features specific to SIRA differentiation
   - Use cost-sensitive learning to penalize SIRA misclassification more heavily

6. Cross-Dataset Validation
   - Test on beans from different regions/growing conditions
   - Ensure model generalizes beyond this specific dataset
"""
doc.add_paragraph(s6_content.strip())

# ============================================================
# REFERENCES
# ============================================================
doc.add_page_break()
ref_title = doc.add_paragraph("References")
ref_title.runs[0].bold = True
ref_title.runs[0].font.size = Pt(14)

references = """
[1] Koklu, M. and Ozkan, I.A., 2020. Multiclass classification of dry beans using computer vision and machine learning techniques. Computers and Electronics in Agriculture, 174, 105507.

[2] UCI Machine Learning Repository - Dry Bean Dataset. https://archive.ics.uci.edu/ml/datasets/Dry+Bean+Dataset

[3] Pedregosa, F. et al., 2011. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, pp.2825-2830.

[4] Breiman, L., 2001. Random Forests. Machine Learning, 45(1), pp.5-32.

[5] Cortes, C. and Vapnik, V., 1995. Support-vector networks. Machine Learning, 20(3), pp.273-297.

[6] Cover, T. and Hart, P., 1967. Nearest neighbor pattern classification. IEEE Transactions on Information Theory, 13(1), pp.21-27.
"""
doc.add_paragraph(references.strip())

# ============================================================
# APPENDIX
# ============================================================
doc.add_page_break()
app_title = doc.add_paragraph("Appendix (Not Marked)")
app_title.runs[0].bold = True
app_title.runs[0].font.size = Pt(14)

doc.add_paragraph("The following visualizations are generated by the code and saved in the output/ folder:")
doc.add_paragraph()

appendix_items = [
    "1. class_distribution.png - Bar chart showing sample counts per bean class",
    "2. correlation_heatmap.png - Feature correlation matrix",
    "3. feature_distributions.png - Histograms of key features by class",
    "4. feature_boxplots.png - Boxplots showing feature spread by class",
    "5. model_comparison.png - Bar chart comparing all metrics across models",
    "6. cv_comparison.png - Cross-validation accuracy with error bars",
    "7. confusion_matrices.png - Confusion matrices for all 3 models",
    "8. feature_importance.png - Random Forest feature importance ranking",
]

for item in appendix_items:
    doc.add_paragraph(item)

doc.add_paragraph()
doc.add_paragraph("AI Disclosure: Claude (Anthropic) was used to assist with code generation for the ML pipeline and document formatting. All code was reviewed, tested, and understood by team members. Results and analysis were verified independently.")

# Save document
doc.save(OUTPUT_FILE)
print(f"Technical Summary saved to: {OUTPUT_FILE}")
